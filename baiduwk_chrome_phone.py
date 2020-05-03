'''
import time
import os
from selenium import webdriver
options = webdriver.FirefoxOptions()
options.add_argument('Mozilla/5.0 (Android 9.0; Mobile; rv:63.0) Gecko/63.0 Firefox/63.0')
driver = webdriver.Firefox(firefox_options=options)

driver.get('https://www.baidu.com')
driver.get('https://wenku.baidu.com/view/81262400376baf1ffd4fad26.html')
'''

#爬到的数据不完整
def baiduwenkudocx(links):
    from selenium import webdriver
    import docx,time

    options = webdriver.ChromeOptions()
    #options.add_argument('user-agent="Mozilla/5.0 (Android 9.0; Mobile; rv:63.0) Gecko/63.0 Firefox/63.0"')
    options.add_argument('user-agent="Mozilla/5.0 (Linux; Android 4.0.4; Galaxy Nexus Build/IMM76B) AppleWebKit/535.19 (KHTML, like Gecko) Chrome/18.0.1025.133 Mobile Safari/535.19"')
    driver=webdriver.Chrome(chrome_options=options)
    driver.set_window_size(1500,800)
    #driver.get('https://www.baidu.com')


    classname=['txt']
    xpathname=['//p[@class]','//span[@class]']
    i=1
    documents = ['',]

    for link in links:
        documents.append('')
        documents[i] = docx.Document()
        n = 0
        driver.get(link)
        clickcomment=driver.find_element_by_class_name('foldpagewg-text')
        clickcomment.click()
        time.sleep(30)
        # clickcomment=driver.find_element_by_xpath('//span[@class="pagerwg-arrow-lower"]')
        # if clickcomment:
        #     clickcomment.click()
        # #selenium.common.exceptions.WebDriverException: Message: unknown error: Element <span class="pagerwg-arrow-lower"></span> is not clickable
        #

        comments=driver.find_elements_by_class_name(classname[0])
        # comments=driver.find_elements_by_xpath(xpathname[0])

        for eachcomment in comments:
            if eachcomment.text=='打开文库App，查看更多同类文档' or eachcomment.text=='下载原文档，方便随时阅读':
                break
            if eachcomment.text=='百度文库':
                continue
            n=n+1
            print(eachcomment.text)
            #s=eachcomment.text
            #s=str('        '+s)
            #document.add_paragraph('\n')
            documents[i].add_paragraph(eachcomment.text)#写入段落
        documents[i].save("test%d.docx"%i)  # 保存文档
        i+=1
        if n>1:
            print('\n%d comments'%n)
        else:
            print('\n%d comment'%n)


links=['https://wenku.baidu.com/view/c5dee18eed630b1c59eeb5dc.html'
          ]
baiduwenkudocx(links)